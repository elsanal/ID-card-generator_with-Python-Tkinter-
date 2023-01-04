from docx.enum.text import WD_BREAK
from docx import Document
import os
from tkinter import messagebox
import sex
from errors import empty_input, photo_not_selected
from print_image import image_printer
from count_generator import numGenerator
import validation
from qrCode_generator import qr_code_generator


def printer(name, surname, birth, birth_place, tel, email, resid, add, cp, DL, DE, filename, card_num, doc_path,
            is_printing):
    print(sex.gender)
    data = 'Numéro de carte (卡号) : ' + str(card_num) + '\n\n' + \
           'Nom (性) : ' + str(name) + '\n\n' + \
           'Prénom (名) : ' + str(surname) + '\n\n' + \
           'Date de naissance (出日) : ' + str(birth) + ' A ' + str(birth_place) + '\n\n' + \
           'Sexe (性别) : ' + str(sex.gender) + '\n\n' + \
           'Téléphone (手机号) : ' + str(tel) + '\n\n' + \
           'Email (邮件) : ' + str(email) + '\n\n' + \
           'Résidence (住宅) : ' + str(resid) + '\n\n' + \
           'Addresse (地址) : ' + str(add) + '\n\n' + \
           'Date de livraison (颁发时间) : ' + str(DL) + '\n\n' + \
           "Date d'expiration (到期时间) : " + str(DE)
    inputs = [ {'key': 'Nom', 'value':name, 'important':True}, {'key': 'Prénom', 'value':surname, 'important':True},
               {'key': 'Date de naissance', 'value':birth, 'important':True},{'key': 'Lieu de naissance', 'value': birth_place, 'important':True},
               {'key': 'Sexe', 'value':sex.gender, 'important':True}, {'key': 'Numero de telephone', 'value':tel, 'important':False},
               {'key': 'Email', 'value':email, 'important':False},{'key': 'Residence', 'value':resid, 'important':False},
               {'key': 'Adresse', 'value':add, 'important':False},{'key': 'Code postal', 'value':cp, 'important':False},
               {'key': 'Date de livraison', 'value': DL, 'important':True},{'key': "Date d'expiration", 'value':DE, 'important':True},
               ]

    if not empty_input(inputs) and not photo_not_selected(filename):
        try:
            curr_path = name + ' ' + surname + ' N°ABFPK ' + card_num + '/'
            path = os.path.join(doc_path, curr_path)
            os.mkdir(path)
        except:
            curr_path = name + ' ' + surname + ' N°ABFPK ' + card_num + '/'
            path = str(doc_path) + '/' + str(curr_path)

        qr_code_generator(data, path)
        ##################### Infos document ###############################
        info_doc = Document()
        info_doc.add_heading("Ambassade du Burkina Faso A Pekin \n\t\t\t 布基纳法索大使馆", level=0)
        info_doc.add_heading("\t\t\t\tCarte consulaire", level=1)
        para = info_doc.add_paragraph()
        para_run = para.add_run()
        para_run.add_text('Numéro de carte : ABFPK/' + str(card_num))
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Nom : ' + name)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Prénom : ' + surname)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Date de naissance : ' + birth + ' A ' + birth_place)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Sexe : ' + sex.gender)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Téléphone : ' + tel)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Email : ' + email)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Résidence : ' + resid)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Adresse : ' + add)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Code postal : ' + cp)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text('Date de livraison : ' + DL)
        para_run.add_break(break_type=WD_BREAK.LINE)
        para_run.add_text("Date d'expiration : " + DE)
        para_run.add_break(break_type=WD_BREAK.LINE)

        info_doc.save(path + 'information.docx')
        validation.show_card(card_num, name, surname, birth, birth_place, add, DL, DE, filename, path, is_save=True)
        numGenerator('../Data/card_numb.txt', card_num, is_to_add=True)
        messagebox.showinfo('Information', f"Les informations ont été bien sauvegardé dans\n{path}")
        if is_printing:
            front = open(file=f"{path}page_devant.png", mode='r', encoding="utf8")
            back = open(file=f"{path}page_arriere.png", mode='r', encoding="utf8")
            try:
                image_printer(front)
                image_printer(back)
                messagebox.showinfo("Erreur", "La carte est imprimée")

            except:
                messagebox.showerror("Erreur", "Connectez votre ordinateur a un imprimante")
